import os
import sys
import ast
import io
import openai
import json
import docx
import glob
import docx2txt
import logging
import azure.functions as func 
import pandas as pd
from azure.identity import DefaultAzureCredential

from tqdm import tqdm 
from pprint import pprint
from tqdm import tqdm
from pathlib import Path

from langchain.embeddings import OpenAIEmbeddings
from azure.storage.blob import BlobServiceClient
from azure.core.credentials import AzureKeyCredential  
from azure.search.documents import SearchClient  

def clean_header(header_list):
    new_header = []
    last_element = header_list[-1].lower()
    for i in range(len(header_list)-1):
        if header_list[i].lower()!=last_element:
            new_header.append(header_list[i])
    new_header.append(header_list[-1])
    return ", ".join(new_header)

def clean_list_of_elements(list_of_elements):
    elt = list_of_elements[0]
    cleaned_list_of_elements = [elt]
    cpt_test = 1
    i = 1
    while i<len(list_of_elements):
        if list_of_elements[i] != elt:
            cleaned_list_of_elements.append(list_of_elements[i])
            elt = list_of_elements[i]
        i+=1
    return cleaned_list_of_elements

def split_long_blocs(text_in, max_size_bloc, overlap_size=200):
    
    stop_bloc_chars = [".","?","\n"]
    words_text_in = text_in.split()
    
    cpt_blocs = 0
    i=0
    blocs = [[]]
    end = False
    while i < len(words_text_in) and end == False:
        if len(blocs[cpt_blocs]) < max_size_bloc:
            blocs[cpt_blocs].append(words_text_in[i])
            i+=1
        else:
            for c in stop_bloc_chars:
                if c in words_text_in[i] or i==len(words_text_in)-1:
                    blocs[cpt_blocs].append(words_text_in[i])
                    blocs.append([])
                    cpt_blocs+=1
                    i -= overlap_size
                    break
            i+=1
    for i in range(len(blocs)):
        blocs[i] = " ".join(blocs[i])
    return blocs

def generate_blocs(list_of_elements_input, header_text, path,  max_size_bloc=200):
    list_of_elements = clean_list_of_elements(list_of_elements_input)
    texts = []
    metadatas = []
    i=0
    cpt = 1
    bloc_header = ""
    path_metatdata = path.split('/')[-1]
    while True:
        if i>=len(list_of_elements)-1:
            break
        if list_of_elements[i].replace('\n','')=="":
            bloc_header = list_of_elements[i+1]
            i+=2
        else: 
            text_bloc = "\n".join([list_of_elements[i], list_of_elements[i+1]])
            if len(text_bloc.split()) > max_size_bloc:
                splited_texts = split_long_blocs(text_bloc, max_size_bloc, overlap_size=100)
                for cpt_sub_bloc, text_bloc in enumerate(splited_texts):
                    text_to_append = "\n".join([header_text, bloc_header, text_bloc]).replace('\n\n','\n')
                    texts.append(text_to_append)
                    metadatas.append({"source":path_metatdata+"_Bloc-"+str(cpt)+"-"+str(cpt_sub_bloc+1)})
            else:
                text_to_append = "\n".join([header_text, bloc_header, list_of_elements[i], list_of_elements[i+1]]).replace('\n\n','\n')
                texts.append(text_to_append)
                metadatas.append({"source":path_metatdata+"_Bloc-"+str(cpt)})
            cpt+=1
            i+=2
    return texts, metadatas

def split_docx_with_tables_data(path, doc_file=None):
    ls = []
    header = []
    texts = []
    metadatas = []
    title = ""
    if doc_file:
        doc = docx.Document(doc_file)
    else:
        doc = docx.Document(path)
    nbr_tables = len(doc.tables)

    tables = doc.tables[0]
    for row in tables.rows:
        for cell in row.cells:
            header.append(" ".join([paragraph.text for paragraph in cell.paragraphs]))
    header = clean_header(header)
    
    for i in range(1, nbr_tables):
        ls = []
        tables = doc.tables[i]
        for row in tables.rows:
            for cell in row.cells:
                text = "\n".join([paragraph.text for paragraph in cell.paragraphs])
                ls.append(text.replace('\n\n','\n'))
        texts_tmp, metadatas_tmp = generate_blocs(ls, header, path)
        texts += texts_tmp
        metadatas += metadatas_tmp        
    return texts, metadatas

def process_output(result, sources):
    if "Aucun résultat" in result['output_text']:
        print("\n")
        print("#"*100)
        print("Aucun résultat.")
        print("Soyez plus précis dans votre demande.")
        print("#"*100)
        print("\n\n")
        return 
    
    tmp = result['output_text'].split('SOURCES:')
    if '\n' in tmp[0]:
        answer = tmp[0].split('\n')
    else:
        answer = [tmp[0]]
    
    sources_out = []
    for s in sources:
        s = s.metadata['source'].split(".")[0]
        if s not in sources_out:
            sources_out.append(s)

    print("\n")
    print("#"*100)
    print("RESULTAT  : ")
    for a in answer:
        print(a)
    print("*"*100)
    print("SOURCE : ")
    for s in sources_out:
        print(s)
    print("#"*100)

def generate_embeddings(text):
    response = openai.Embedding.create(
        input=text, engine="")
    embeddings = response['data'][0]['embedding']
    return embeddings

def main(inputblob: func.InputStream):
    logging.info(f"Python blob trigger function processed blob \n"
                 f"Name: {inputblob.name}\n"
                 f"Blob Size: {inputblob.length} bytes\n"
                 f"Blob URI: {inputblob.uri} bytes")

    #TEST IF DOCX OR NOT
    if(not inputblob.name.endswith(".docx")):
        logging.info(f"{inputblob.name} isn't a DOCX file, aborting.")
        return
    

    # Gather all Functions parameters
    OPENAI_API_TYPE = os.environ["OPENAI_API_TYPE"]
    OPENAI_API_BASE = os.environ["OPENAI_API_BASE"]
    OPENAI_API_VERSION = os.environ["OPENAI_API_VERSION"]
    OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]
    OPENAI_DEPLOYMENT_NAME = os.environ["OPENAI_DEPLOYMENT_NAME"]
    OPENAI_EMBEDDED_DEPLOYMENT_NAME = os.environ["OPENAI_EMBEDDED_DEPLOYMENT_NAME"]
    OPENAI_EMBEDDED_CHUNK_SIZE= os.environ["OPENAI_EMBEDDED_CHUNK_SIZE"] 
    SAVE_INDEX = os.environ["SAVE_INDEX"] 
    AZURE_STORAGE_URI = os.environ["AZURE_STORAGE_URI"]
    AZURE_STORAGE_CONTAINER = os.environ["AZURE_STORAGE_CONTAINER"]
    SEARCH_SERVICE_ENDPOINT = os.environ["SEARCH_SERVICE_ENDPOINT"]
    SEARCH_INDEX_NAME = os.environ["SEARCH_INDEX_NAME"]
    SEARCH_KEY = os.environ["SEARCH_KEY"]
    AZURE_STORAGE_SAS = os.environ["AZURE_STORAGE_SAS"]

    Searchcredential = AzureKeyCredential(SEARCH_KEY)
    azure_credential = AZURE_STORAGE_SAS
    
    search_client = SearchClient(endpoint=SEARCH_SERVICE_ENDPOINT, index_name=SEARCH_INDEX_NAME, credential=Searchcredential)
    
    blob_service_client = BlobServiceClient(
        account_url=AZURE_STORAGE_URI, 
        credential=azure_credential)
    
    blobName =  inputblob.name.rsplit('/', 1)[-1]
    blobURIWithoutContainer = inputblob.name.replace(AZURE_STORAGE_CONTAINER,"")
    blobWithoutExtension =  blobName.replace(".DOCX","")
    blobWithoutExtension =  blobWithoutExtension.replace(".docx","")

    blob_container = blob_service_client.get_container_client(AZURE_STORAGE_CONTAINER)
    
    blob = blob_container.get_blob_client(blob=blobURIWithoutContainer)

    logging.info("Creating embedding model")
    embeddings = OpenAIEmbeddings(
        deployment=OPENAI_EMBEDDED_DEPLOYMENT_NAME,
        openai_api_base=OPENAI_API_BASE,
        openai_api_type=OPENAI_API_TYPE,
        openai_api_version=OPENAI_API_VERSION,
        openai_api_key=OPENAI_API_KEY,
        chunk_size=OPENAI_EMBEDDED_CHUNK_SIZE
    )
    if SAVE_INDEX:
        logging.info("Load texts & metadatas")

        logging.info(f"Loading {inputblob.name} document")
        all_chunk_dicts = []
        
        #Store locally the file
        doc = open(blobName, "wb")
        doc.write(blob.download_blob().content_as_bytes())
        doc.close()

        
        logging.info(f"Local copy done")        
       
        all_chunk_dicts = []

        text, metadata = split_docx_with_tables_data(blobName, doc_file=None)
        for idx, (chunk_text, chunk_metadata) in enumerate(zip(text, metadata)):
            chunk_dict = {
                "text": chunk_text,
                "source": chunk_metadata["source"],
                "docname": inputblob.name,
                "id": f"{blobWithoutExtension}_{idx}"
            }
            all_chunk_dicts.append(chunk_dict)

        logging.info("Building chunks dataframe")    

        df_chunks = pd.DataFrame(all_chunk_dicts)

        logging.info("Saving chunks dataframe")
        logging.info(df_chunks)
        df_chunks.to_json(f"{blobWithoutExtension}_chunks.json", orient="records", lines=False, force_ascii=False)
        logging.info("Uploading chunks to datastore")
        
        with open(file=f"{blobWithoutExtension}_chunks.json", mode="rb") as data:
            blob_container.upload_blob(name =f"/tmp/{blobWithoutExtension}_chunks.json", data=data, overwrite=True)
            
        for item in tqdm(df_chunks):
            #title = item['title']
            content = item['text']
            #title_embeddings = generate_embeddings(title)
            content_embeddings = generate_embeddings(content)
            #item['titleVector'] = title_embeddings
            item['contentVector'] = content_embeddings
            item['@search.action'] = 'upload'

        # Output embeddings to docVectors.json file
        df_chunks.to_json(f"{blobWithoutExtension}_docVectors.json", orient="records", lines=False, force_ascii=False)
        logging.info("Uploading Vectors to datastore")

        with open(file=f"{blobWithoutExtension}_docVectors.json", mode="rb") as data:
            blob_container.upload_blob(name =f"/tmp/{blobWithoutExtension}_docVectors.json", data=data, overwrite=True)
        
        # documents = json.load(data)
        result = search_client.upload_documents(f"{blobWithoutExtension}_docVectors.json") 
        logging.info(f"Uploaded documents to search done") 

    else:
    # if exists load 
         print("Not Search")
    
    return "Done"